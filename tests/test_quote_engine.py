"""Tests for quote ingestion behavior."""

from pathlib import Path
import subprocess

import pytest
from docx import Document

from QuoteEngine import (
    CSVIngestor,
    DocxIngestor,
    Ingestor,
    InvalidQuoteFormatError,
    PDFIngestor,
    QuoteModel,
    TextIngestor,
    UnsupportedFileTypeError,
)


BASE_DIR = Path(__file__).resolve().parents[1]


def test_quote_model_representation():
    """QuoteModel prints in the Udacity-required format."""
    quote = QuoteModel("Readable code matters", "Devin")

    assert repr(quote) == '"Readable code matters" - Devin'
    assert str(quote) == '"Readable code matters" - Devin'


def test_text_ingestor_parses_standard_lines(tmp_path):
    """TextIngestor turns valid text lines into QuoteModel objects."""
    quote_file = tmp_path / "quotes.txt"
    quote_file.write_text('"Line one" - Author One\n', encoding="utf-8")

    quotes = TextIngestor.parse(quote_file)

    assert len(quotes) == 1
    assert quotes[0].body == "Line one"
    assert quotes[0].author == "Author One"


def test_text_ingestor_rejects_bad_quote_format(tmp_path):
    """Invalid text lines fail with a human-readable custom exception."""
    quote_file = tmp_path / "quotes.txt"
    quote_file.write_text("missing author separator\n", encoding="utf-8")

    with pytest.raises(InvalidQuoteFormatError, match="Expected the format"):
        TextIngestor.parse(quote_file)


def test_text_ingestor_accepts_utf8_bom(tmp_path):
    """UTF-8 BOM markers are removed without changing quote content."""
    quote_file = tmp_path / "bom-quotes.txt"
    quote_file.write_text(
        '"A clean beginning" - Devin\n',
        encoding="utf-8-sig",
    )

    quotes = TextIngestor.parse(quote_file)

    assert quotes == [QuoteModel("A clean beginning", "Devin")]


def test_csv_ingestor_normalizes_padded_uppercase_columns(tmp_path):
    """CSV column matching remains whitespace and case insensitive."""
    quote_file = tmp_path / "quotes.csv"
    quote_file.write_text(
        ' BODY , AUTHOR \n"Useful defaults","Devin"\n',
        encoding="utf-8",
    )

    quotes = CSVIngestor.parse(quote_file)

    assert quotes == [QuoteModel("Useful defaults", "Devin")]


def test_docx_ingestor_skips_blanks_and_rejects_invalid_lines(tmp_path):
    """DOCX spacing is ignored while malformed quote lines still fail."""
    quote_file = tmp_path / "quotes.docx"
    document = Document()
    document.add_paragraph("")
    document.add_paragraph('"Valid quote" - Devin')
    document.add_paragraph("invalid quote line")
    document.save(quote_file)

    with pytest.raises(InvalidQuoteFormatError, match="Expected the format"):
        DocxIngestor.parse(quote_file)


def test_pdf_ingestor_uses_configured_pdftotext(monkeypatch, tmp_path):
    """PDFTOTEXT_PATH selects the subprocess parser branch directly."""
    pdf_path = tmp_path / "quotes.pdf"
    pdf_path.write_bytes(b"placeholder")
    converted_path = None

    def fake_run(command, **options):
        nonlocal converted_path
        assert command[:2] == ["custom-pdftotext", str(pdf_path)]
        assert options == {
            "check": True,
            "capture_output": True,
            "text": True,
        }
        converted_path = Path(command[2])
        converted_path.write_text(
            '"Subprocesses need tests" - Devin\n',
            encoding="utf-8",
        )
        return subprocess.CompletedProcess(command, returncode=0)

    monkeypatch.setenv("PDFTOTEXT_PATH", "custom-pdftotext")
    monkeypatch.setattr(subprocess, "run", fake_run)

    quotes = PDFIngestor.parse(pdf_path)

    assert quotes == [QuoteModel("Subprocesses need tests", "Devin")]
    assert converted_path is not None
    assert not converted_path.exists()


@pytest.mark.parametrize(
    "fixture_path,expected_count",
    [
        ("_data/SimpleLines/SimpleLines.csv", 5),
        ("_data/SimpleLines/SimpleLines.docx", 5),
        ("_data/SimpleLines/SimpleLines.pdf", 5),
        ("_data/DogQuotes/DogQuotesTXT.txt", 2),
    ],
)
def test_ingestor_parses_supported_fixture_types(fixture_path, expected_count):
    """The facade parses the CSV, DOCX, PDF, and TXT fixture formats."""
    quotes = Ingestor.parse(BASE_DIR / fixture_path)

    assert len(quotes) >= expected_count
    assert all(isinstance(quote, QuoteModel) for quote in quotes)


def test_ingestor_discovers_supported_files_in_directory(tmp_path):
    """Directory parsing uses os.walk to collect supported quote files."""
    nested = tmp_path / "nested"
    nested.mkdir()
    (tmp_path / "first.txt").write_text('"One" - A\n', encoding="utf-8")
    (nested / "second.txt").write_text('"Two" - B\n', encoding="utf-8")
    (nested / "ignore.md").write_text('"Nope" - C\n', encoding="utf-8")

    quotes = Ingestor.parse_directory(tmp_path)

    assert [quote.body for quote in quotes] == ["One", "Two"]


def test_ingestor_rejects_unsupported_file_type(tmp_path):
    """Unsupported files raise a custom exception."""
    unsupported = tmp_path / "quotes.md"
    unsupported.write_text('"Nope" - C\n', encoding="utf-8")

    with pytest.raises(UnsupportedFileTypeError):
        Ingestor.parse(unsupported)
